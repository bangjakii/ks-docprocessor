"""
KS Pinecone Indexer — Waralalo Group
====================================
Ambil dokumen yang sudah ditata oleh ingest_archive.py (--apply) dari folder
hasil (default D:\\Arsip_Rapih), ekstrak teks (pdfplumber + OCR untuk scan),
pecah PER-HALAMAN, lalu index ke Pinecone dengan INTEGRATED EMBEDDING
(Pinecone yang embed server-side — tidak perlu OpenAI).

Granularitas pencarian = chunk per-halaman (bukan pemotongan file fisik). Lihat
keputusan di memory: bundel campur disimpan utuh, granularitas via chunking.

Baca scan: default Claude VISION (teks bersih), bukan Tesseract. Ganti via
env OCR_ENGINE=tesseract kalau mau banding/hemat.

Cara pakai:
    pip install -r requirements.txt
    # (pinecone[asyncio] pdfplumber pdf2image pytesseract pillow python-docx openpyxl anthropic python-dotenv)
    python index_to_pinecone.py                     # index seluruh Arsip_Rapih (vision)
    python index_to_pinecone.py --limit 50           # uji 50 file dulu
    python index_to_pinecone.py --query "faktur pajak PT PAL"   # cari (tes)

Env (.env): PINECONE_API_KEY, ANTHROPIC_API_KEY (vision), POPPLER_PATH (render),
            TESSERACT_PATH (kalau OCR_ENGINE=tesseract), VISION_MODEL (opsional)
"""

import os
import io
import re
import sys
import json
import time
import base64
import hashlib
import argparse
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from datetime import datetime

# Konsol Windows default cp1252 → emoji/box-drawing di print() bisa crash.
# Paksa stdout/stderr ke utf-8 utk semua entry point yg meng-import modul ini.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except Exception:
        pass

import pdfplumber
from dotenv import load_dotenv

load_dotenv()

# ── Konfigurasi default ───────────────────────────────────────────────────────
DEST_DIR      = Path(os.getenv("INGEST_DEST", r"D:\Arsip_Rapih"))
INDEX_NAME    = "ks-documents"
EMBED_MODEL   = "multilingual-e5-large"   # integrated, 1024-dim, multilingual (ID)
NAMESPACE     = "__default__"
CHUNK_SIZE    = 1500     # karakter per chunk dalam satu halaman
CHUNK_OVERLAP = 200
# OCR_PAGE_CAP didefinisikan di blok OCR di bawah (default/high/low per doc-type).
BATCH         = 96       # upsert_records per batch
# Paralelisasi: OCR per-file diproses paralel (CPU tesseract/render + I/O vision).
# Default ~ 8 core + sedikit over utk nutup tunggu I/O. >16 = thrash + rate-limit.
WORKERS       = int(os.getenv("INDEX_WORKERS", str(min(12, (os.cpu_count() or 8) + 4))))
CHECKPOINT    = Path("pinecone_indexed.json")   # set base_id yang sudah selesai

# Metadata field yang dibawa per chunk (selain _id & chunk_text yang di-embed).
META_FIELDS = ("doc_name", "company", "counterparty", "department", "project",
               "subfolder", "relpath", "source_file", "expire_date", "doc_number")

# ── Baca scan: HYBRID (default) — Tesseract utk cetak, Vision utk tulisan tangan ─
# Tesseract jalan duluan (gratis); confidence-nya jadi router: cetak rapi → pakai
# tesseract, tulisan tangan/stempel/scan ancur → eskalasi ke Claude vision.
# Halaman gambar/kosong minim-teks → di-SKIP (biar vision ga ngarang).
# Override: OCR_ENGINE=tesseract (paksa tesseract) | vision (paksa vision) | hybrid.
OCR_ENGINE     = os.getenv("OCR_ENGINE", "hybrid").lower()     # "hybrid" | "vision" | "tesseract"
# Default HEMAT: haiku utk vision normal, sonnet utk scan TERJELEK (conf rendah).
# Naikkan ke opus (VISION_MODEL_STRONG=claude-opus-4-8) kalau mau kualitas maksimal.
VISION_MODEL        = os.getenv("VISION_MODEL", "claude-haiku-4-5")
VISION_MODEL_STRONG = os.getenv("VISION_MODEL_STRONG", "claude-sonnet-4-6")
# Knob router hybrid (bisa dikalibrasi via env):
OCR_CONF_OK     = float(os.getenv("OCR_CONF_OK", "75"))      # conf >= ini & rapi → keep tesseract (gratis)
OCR_CONF_STRONG = float(os.getenv("OCR_CONF_STRONG", "55"))  # conf < ini → vision pakai model KUAT
OCR_MIN_INK     = int(os.getenv("OCR_MIN_INK", "4"))         # kata < ini & low-conf → blank/drawing → skip
# OCR_PAGE_CAP per DOC-TYPE (B): bundel Finance/invoice TINGGI (faktur nyempil di dalam),
# engineering/drawing RENDAH (hal 2+ cuma gambar). Lainnya = default.
OCR_PAGE_CAP   = int(os.getenv("OCR_PAGE_CAP", "5"))         # default/medium
OCR_CAP_HIGH   = int(os.getenv("OCR_CAP_HIGH", "8"))         # Finance/invoice/faktur/pinjaman (faktur ≤hal4 + kwitansi/bukti)
OCR_CAP_LOW    = int(os.getenv("OCR_CAP_LOW", "2"))          # engineering/drawing/brosur
_CAP_HIGH_KW = ("invoice", "faktur", "tagihan", "pembayaran", "kwitansi", "finance",
                "keuangan", "pinjaman", "borrower", "bundel")
_CAP_LOW_KW  = ("engineering", "drawing", "gambar", "general arrangement", " ga ", "abs",
                "brosur", "spesifikasi", "katalog", "datasheet")


def cap_for(meta: dict) -> int:
    """OCR_PAGE_CAP sesuai doc-type (dari department/subfolder/doc_name)."""
    blob = " " + re.sub(r"[^a-z0-9 ]", " ",
                        " ".join(str(meta.get(x, "")) for x in
                                 ("department", "subfolder", "doc_name")).lower()) + " "
    if any(k in blob for k in _CAP_LOW_KW):
        return OCR_CAP_LOW
    if any(k in blob for k in _CAP_HIGH_KW):
        return OCR_CAP_HIGH
    return OCR_PAGE_CAP
# Ekstraksi field terstruktur (doc_number/expire_date/counterparty): regex gratis +
# LLM fallback tertarget. Matikan dgn EXTRACT_FIELDS=0 (atau LLM saja: EXTRACT_LLM=0).
ENABLE_FIELD_EXTRACT = os.getenv("EXTRACT_FIELDS", "1").lower() not in ("0", "false", "no")
EXTRACT_USE_LLM      = os.getenv("EXTRACT_LLM", "1").lower() not in ("0", "false", "no")
TESSERACT_PATH = os.getenv("TESSERACT_PATH")
POPPLER_PATH   = os.getenv("POPPLER_PATH")
try:
    import pytesseract
    from pdf2image import convert_from_path, pdfinfo_from_path
    if TESSERACT_PATH:
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    _RENDER_OK = True
except ImportError:
    _RENDER_OK = False


def _poppler_pagecount(pdf_path: Path) -> int:
    """Jumlah halaman menurut poppler. Andalan saat pdfplumber lapor 0 halaman
    (sering terjadi di PDF scan dgn page-tree yg tak terbaca pdfminer)."""
    if not _RENDER_OK:
        return 0
    try:
        info = pdfinfo_from_path(str(pdf_path), poppler_path=POPPLER_PATH or None)
        return int(info.get("Pages", 0))
    except Exception:
        return 0

_NOTEXT = "[TIDAK TERBACA]"
_VISION_PROMPT = (
    "Transkripsikan SELURUH teks pada halaman dokumen ini secara verbatim (persis apa "
    "adanya) dalam bahasa aslinya: kop surat, nomor & tanggal dokumen, nama perusahaan, "
    "angka, dan isi tabel (tulis baris per baris). Pertahankan urutan baca. JANGAN "
    "merangkum, menerjemahkan, menambah komentar, atau memakai format markdown "
    "(tanpa **, #, atau bullet). Keluarkan LANGSUNG teks transkripsinya saja TANPA "
    "kalimat pembuka apa pun (jangan tulis 'Berikut', 'Saya akan', 'Teks pada dokumen', dll). "
    "PENTING: transkripsikan HANYA teks yang benar-benar terlihat. JANGAN PERNAH menebak, "
    "mengarang, atau melengkapi nama perusahaan, nomor, tanggal, atau isi yang tidak jelas. "
    "Kalau ini gambar teknik/denah/diagram, tetap transkrip teks yang ada (kop, nama kapal, "
    "dimensi, skala, nomor & judul gambar, label). "
    f"Kalau halaman benar-benar TIDAK ADA teks atau tak terbaca, balas HANYA token {_NOTEXT} "
    "— JANGAN menulis penjelasan, alasan, atau permintaan maaf apa pun."
)

# Pola kalimat pembuka yang kadang bocor dari model meski sudah dilarang.
_PREAMBLE_RE = re.compile(
    r"^\s*(?:berikut(?:\s+adalah)?|saya\s+akan|teks\s+(?:yang\s+)?(?:terlihat|pada)|"
    r"hasil\s+transkrip\w*|transkrip\w*)[^\n:]*:\s*", re.IGNORECASE)
# Pola PENOLAKAN/permintaan maaf model (tak terbaca tapi tanpa sentinel) — jangan di-index.
_REFUSAL_RE = re.compile(
    r"(saya\s+(?:tidak|tak|ga|gak)\s+(?:dapat|bisa|mampu)|tidak\s+dapat\s+(?:men|membaca)|"
    r"maaf,?\s+(?:saya|tidak)|kualitas\s+(?:gambar|citra|dokumen)[^.\n]{0,40}"
    r"(?:rendah|buruk|kurang)|tidak\s+ada\s+teks\s+(?:yang\s+)?(?:terlihat|terbaca|jelas)|"
    r"i('?m| am)?\s*(?:cannot|can't|am\s+unable|unable\s+to))", re.IGNORECASE)


def _clean_vision(t: str) -> str:
    """Buang preamble, sentinel, & kalimat penolakan supaya tidak meracuni index."""
    t = (t or "").strip()
    t = _PREAMBLE_RE.sub("", t).strip()
    if not t or t.upper().startswith(_NOTEXT) or t.upper() == _NOTEXT.strip("[]"):
        return ""
    if _REFUSAL_RE.search(t[:120]):       # penolakan selalu di awal respons
        return ""
    return t

_claude_client = None
def _get_claude():
    global _claude_client
    if _claude_client is None:
        from anthropic import Anthropic
        _claude_client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    return _claude_client


def _render_page(pdf_path: Path, page_no: int, dpi: int = 150):
    """Render satu halaman PDF → PIL image (1-indexed). None kalau gagal."""
    if not _RENDER_OK:
        return None
    try:
        imgs = convert_from_path(str(pdf_path), dpi=dpi, first_page=page_no,
                                 last_page=page_no, poppler_path=POPPLER_PATH or None)
        return imgs[0] if imgs else None
    except Exception:
        return None


def _render_range(pdf_path: Path, last: int, dpi: int = 180):
    """Render halaman 1..last dalam SATU panggilan poppler → {page_no: PIL image}.
    Jauh lebih cepat dari render per-halaman (poppler cuma spawn 1×, bukan N×)."""
    if not _RENDER_OK or last < 1:
        return {}
    try:
        imgs = convert_from_path(str(pdf_path), dpi=dpi, first_page=1,
                                 last_page=last, poppler_path=POPPLER_PATH or None)
        return {i + 1: im for i, im in enumerate(imgs)}
    except Exception:
        return {}


def _tesseract_scored(img):
    """OCR + sinyal routing dalam SATU pass (gratis). Kembalikan (text, n_kata, conf)."""
    try:
        data = pytesseract.image_to_data(
            img, lang="ind+eng", output_type=pytesseract.Output.DICT)
    except Exception:
        return "", 0, 0.0
    words, confs = [], []
    for txt, c in zip(data["text"], data["conf"]):
        t = (txt or "").strip()
        try:
            c = float(c)
        except (TypeError, ValueError):
            c = -1.0
        if t and c >= 0:
            words.append(t)
            confs.append(c)
    text = " ".join(words)
    conf = sum(confs) / len(confs) if confs else 0.0
    return text, len(words), conf


def _tesseract_page(pdf_path: Path, page_no: int) -> str:
    img = _render_page(pdf_path, page_no, dpi=180)
    if img is None:
        return ""
    return _tesseract_scored(img)[0]


def _vision_image(img, model: str = None) -> str:
    """Kirim PIL image ke Claude vision → teks bersih. Dipisah biar bisa reuse
    render yang sama dari router hybrid (ga render 2x). model=None → VISION_MODEL."""
    if img is None:
        return ""
    try:
        # Downscale: API tolak sisi >8000px; >~2200 juga mubazir (vision di-resize internal).
        # Sekalian payload lebih kecil → lebih cepat & murah.
        if img.width > 2200 or img.height > 2200:
            img = img.copy()
            img.thumbnail((2200, 2200))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()
        msg = _get_claude().messages.create(
            model=model or VISION_MODEL, max_tokens=2048,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64",
                 "media_type": "image/jpeg", "data": b64}},
                {"type": "text", "text": _VISION_PROMPT},
            ]}],
        )
        raw = "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")
        return _clean_vision(raw)
    except Exception as e:
        print(f"    ⚠ vision gagal: {e}")
        return ""


def _vision_page(pdf_path: Path, page_no: int) -> str:
    return _vision_image(_render_page(pdf_path, page_no, dpi=180))


def _ocr_image(img) -> str:
    """Hybrid OCR pada satu PIL image (dipakai PDF maupun file gambar).
    tesseract dulu (gratis); confidence routing → cetak=tesseract,
    tulisan tangan/stempel=vision, gambar/kosong minim-teks=skip."""
    if img is None:
        return ""
    if OCR_ENGINE == "vision":
        return _vision_image(img)
    text, n_words, conf = _tesseract_scored(img)
    if OCR_ENGINE == "tesseract":
        return text
    # ── hybrid ──
    if n_words == 0:
        return ""                                    # ga ada teks → blank/drawing
    if conf >= OCR_CONF_OK and not _looks_garbled(text):
        return text                                  # cetak yakin & rapi → GRATIS ✅
    if n_words < OCR_MIN_INK:
        return ""                                    # low-conf & nyaris kosong → drawing → skip
    # ada tinta tapi ragu → vision. Scan TERJELEK (conf sangat rendah) → model KUAT.
    model = VISION_MODEL_STRONG if conf < OCR_CONF_STRONG else VISION_MODEL
    return _vision_image(img, model=model)


def read_scan_page(pdf_path: Path, page_no: int) -> str:
    """Baca satu halaman scan PDF (render → _ocr_image)."""
    return _ocr_image(_render_page(pdf_path, page_no, dpi=180))


def _looks_garbled(t: str) -> bool:
    """Deteksi layer teks RUSAK (OCR scanner jelek yg ke-embed di PDF): banyak token
    'kata-mash' kepanjangan (HURUF nyambung) atau spasi terlalu sedikit. Teks bersih
    → False. Abaikan deret titik/garis (leader formulir), GUID, & string angka — itu
    bukan teks rusak, biar tidak salah picu re-OCR vision (buang biaya)."""
    t = t.strip()
    if len(t) < 40:
        return False
    t2 = re.sub(r"[._\-–—…·•*=]{4,}", " ", t)        # buang leader titik/garis formulir
    toks = [w for w in t2.split() if w]
    if not toks:
        return True

    def _wordmash(w):                                # huruf nyambung kepanjangan = OCR rusak
        a = sum(c.isalpha() for c in w)
        return len(w) > 18 and a / len(w) > 0.7      # bukan GUID/angka (yg byk non-huruf)

    # ukur kerapatan spasi pada PROSA saja: token panjang non-kata (GUID/kode) dibuang
    prose = re.sub(r"\S{19,}", lambda m: m.group(0) if _wordmash(m.group(0)) else " ", t2)
    long_ratio  = sum(1 for w in toks if _wordmash(w)) / len(toks)
    space_ratio = prose.count(" ") / max(len(prose), 1)
    max_word    = max((len(w) for w in toks if _wordmash(w)), default=0)
    return long_ratio > 0.12 or space_ratio < 0.07 or max_word > 38


def page_texts(pdf_path: Path, cap: int = None) -> list:
    """Kembalikan [(page_number, text)] per halaman. Halaman digital bersih →
    pdfplumber (gratis). Halaman KOSONG (scan) ATAU ber-layer teks RUSAK → dibaca
    ulang via vision (read_scan_page), dibatasi `cap` (default OCR_PAGE_CAP).
    Halaman digital tetap diambil semua; cap HANYA membatasi halaman scan yg di-OCR."""
    if cap is None:
        cap = OCR_PAGE_CAP
    pages = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                pages.append([i + 1, (page.extract_text() or "").strip()])
    except Exception as e:
        print(f"  ⚠ gagal buka {pdf_path.name}: {e}")
        pages = []
    # pdfplumber kadang lapor 0/kurang halaman utk PDF scan tertentu, padahal
    # poppler bisa render. Tambal halaman yg hilang (kosong) → diisi vision di bawah.
    real = _poppler_pagecount(pdf_path)
    for pn in range(len(pages) + 1, real + 1):
        pages.append([pn, ""])
    # Halaman yg perlu OCR (kosong/garbled, dalam cap). Render SEKALI buat semuanya.
    need = [pn for pn, t in pages if pn <= cap and (not t or _looks_garbled(t))]
    if need:
        imgs = _render_range(pdf_path, max(need))         # 1 spawn poppler, bukan N
        for row in pages:
            pn, t = row
            if pn in need:
                v = _ocr_image(imgs.get(pn)).strip()      # hybrid: tesseract→vision
                if v:
                    row[1] = v
    return [(pn, t) for pn, t in pages if t]


# ── Non-PDF: gambar (OCR), Office modern (ekstrak teks) ───────────────────────
_IMG_EXT  = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".gif"}
_DOCX_EXT = {".docx"}
_XLSX_EXT = {".xlsx", ".xlsm"}
# Diindex via FALLBACK record (nama+metadata) saja — legacy/biner tak terbaca teksnya.
_FALLBACK_EXT = {".doc", ".xls", ".ppt", ".pptx", ".rtf", ".dwg", ".dxf", ".zip", ".rar"}
SUPPORTED_EXT = {".pdf"} | _IMG_EXT | _DOCX_EXT | _XLSX_EXT | _FALLBACK_EXT


def image_texts(path: Path) -> list:
    """File gambar (.jpg/.png/…) → 1 'halaman', OCR hybrid (tanpa poppler)."""
    try:
        from PIL import Image
        img = Image.open(path)
        img.load()
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
    except Exception as e:
        print(f"  ⚠ gagal buka gambar {path.name}: {e}")
        return []
    t = _ocr_image(img).strip()
    return [(1, t)] if t else []


def docx_texts(path: Path) -> list:
    """.docx → teks paragraf + tabel (digital, gratis)."""
    try:
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        t = "\n".join(parts).strip()
    except Exception as e:
        print(f"  ⚠ gagal baca docx {path.name}: {e}")
        return []
    return [(1, t)] if t else []


def xlsx_texts(path: Path) -> list:
    """.xlsx/.xlsm → 1 'halaman' per sheet (digital, gratis)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        pages = []
        for i, ws in enumerate(wb.worksheets, 1):
            rows = []
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) for c in row if c is not None and str(c).strip()]
                if vals:
                    rows.append(" | ".join(vals))
            t = "\n".join(rows).strip()
            if t:
                pages.append((i, t))
        wb.close()
        return pages
    except Exception as e:
        print(f"  ⚠ gagal baca xlsx {path.name}: {e}")
        return []


def extract_pages_any(path: Path, cap: int = None) -> list:
    """Dispatcher per ekstensi → [(page, text)]. `cap` membatasi halaman scan PDF
    yg di-OCR (per doc-type via cap_for). Tipe tak terdukung → [] (nanti dapat
    fallback record nama+metadata di records_for)."""
    s = path.suffix.lower()
    if s == ".pdf":
        return page_texts(path, cap=cap)
    if s in _IMG_EXT:
        return image_texts(path)
    if s in _DOCX_EXT:
        return docx_texts(path)
    if s in _XLSX_EXT:
        return xlsx_texts(path)
    return []


# ── Chunking per-halaman ──────────────────────────────────────────────────────

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start += size - overlap
    return chunks


def doc_id(rel: str) -> str:
    """ID stabil per dokumen dari relpath di dalam dest (tahan pindah drive)."""
    return hashlib.md5(rel.encode("utf-8")).hexdigest()[:16]


def _apply_meta(rec: dict, meta: dict) -> dict:
    for f in META_FIELDS:
        v = meta.get(f)
        rec[f] = "" if v is None else str(v)
    return rec


def records_for(pdf_path: Path, base_id: str, meta: dict, pages=None) -> list:
    """Bangun record per (halaman, sub-chunk) dengan metadata file yang utuh.
    `pages` boleh dioper (hasil page_texts yg sudah dihitung) biar tidak OCR 2x."""
    if pages is None:
        pages = page_texts(pdf_path)
    recs, k = [], 0
    for page_no, text in pages:
        for j, chunk in enumerate(chunk_text(text)):
            rec = {"_id": f"{base_id}_p{page_no}_{j}", "chunk_text": chunk,
                   "page_number": page_no, "chunk_index": k,
                   "filename": pdf_path.name}
            recs.append(_apply_meta(rec, meta))
            k += 1
    if not recs:
        # JAMINAN: tiap file minimal 1 record → selalu ke-index & ketemu via
        # nama/metadata walau OCR nol teks (drawing murni / gambar / scan kosong).
        name = meta.get("doc_name") or pdf_path.stem
        ctx = " ".join(str(x) for x in (name, meta.get("project"), meta.get("company"),
                                        meta.get("subfolder")) if x)
        rec = {"_id": f"{base_id}_p0_0", "chunk_text": ctx or name,
               "page_number": 0, "chunk_index": 0, "filename": pdf_path.name}
        recs.append(_apply_meta(rec, meta))
    return recs


# ── Metadata dari archive_log.json ────────────────────────────────────────────

def load_meta_map(dest: Path) -> dict:
    """output_file (absolut) → entri metadata, dari archive_log.json hasil ingest."""
    log = dest / "archive_log.json"
    mp = {}
    if not log.exists():
        print(f"  ⚠ {log} tidak ada — metadata kosong (jalankan ingest --apply dulu).")
        return mp
    try:
        for e in json.loads(log.read_text(encoding="utf-8")):
            if e.get("output_file") and e.get("status") == "ok":
                mp[str(Path(e["output_file"]).resolve())] = e
    except Exception as e:
        print(f"  ⚠ gagal baca log: {e}")
    return mp


# ── Pinecone ──────────────────────────────────────────────────────────────────

def get_pinecone():
    from pinecone import Pinecone
    key = os.getenv("PINECONE_API_KEY")
    if not key:
        print("\n✗ PINECONE_API_KEY tidak ada di .env\n"); raise SystemExit(1)
    return Pinecone(api_key=key)


def retry(fn, tries: int = 5, delay: float = 3.0, what: str = "pinecone"):
    """Ulangi call Pinecone yang gagal (control-plane api.pinecone.io kadang timeout)."""
    last = None
    for i in range(tries):
        try:
            return fn()
        except Exception as e:
            last = e
            if i < tries - 1:
                print(f"    ⏳ {what} gagal ({type(e).__name__}) — coba lagi {i+1}/{tries}...")
                time.sleep(delay)
    raise last


def _index_ready(pc, name: str) -> bool:
    try:
        s = pc.describe_index(name).status
        return bool(s.get("ready") if isinstance(s, dict) else getattr(s, "ready", False))
    except Exception:
        return False


def ensure_index(pc, name: str, model: str):
    if not pc.has_index(name):
        print(f"  📦 Membuat index '{name}' (integrated: {model})...")
        pc.create_index_for_model(
            name=name, cloud="aws", region="us-east-1",
            embed={"model": model, "field_map": {"text": "chunk_text"}},
        )
        for _ in range(120):                 # tunggu maksimum ~2 menit
            if _index_ready(pc, name):
                break
            time.sleep(1)
        print("  ✓ Index siap")
    return pc.Index(name)


def _load_ckpt() -> set:
    if CHECKPOINT.exists():
        try:
            return set(json.loads(CHECKPOINT.read_text(encoding="utf-8")))
        except Exception:
            return set()
    return set()


def _save_ckpt(done: set):
    CHECKPOINT.write_text(json.dumps(sorted(done)), encoding="utf-8")


# ── Meter biaya (token Claude → USD) ──────────────────────────────────────────
_PRICE = {  # $/1M token (input, output)
    "claude-opus-4-8": (5, 25), "claude-opus-4-7": (5, 25), "claude-opus-4-6": (5, 25),
    "claude-sonnet-4-6": (3, 15), "claude-sonnet-4-5": (3, 15), "claude-haiku-4-5": (1, 5),
}
_USAGE = {"in": 0, "out": 0, "cost": 0.0}
_usage_lock = threading.Lock()


def _record_usage(model, usage):
    if not usage:
        return
    i, o = getattr(usage, "input_tokens", 0), getattr(usage, "output_tokens", 0)
    pin, pout = _PRICE.get(model, (3, 15))     # default tier sonnet kalau model tak dikenal
    with _usage_lock:                          # dipanggil dari banyak worker thread
        _USAGE["in"] += i; _USAGE["out"] += o
        _USAGE["cost"] += i / 1e6 * pin + o / 1e6 * pout


def _wrap_client(client):
    if client is None or getattr(client, "_ks_metered", False):
        return client
    orig = client.messages.create
    def create(**kw):
        r = orig(**kw)
        _record_usage(kw.get("model"), getattr(r, "usage", None))
        return r
    client.messages.create = create
    client._ks_metered = True
    return client


def _enable_cost_meter():
    """Bungkus client vision (IX) & extraction (EF) supaya token tiap panggilan
    Claude terakumulasi → biaya bisa dipantau live saat produksi."""
    _wrap_client(_get_claude())
    try:
        import extract_fields as EF
        EF._client = _wrap_client(EF._claude())
    except Exception:
        pass


# ── Indexing utama ────────────────────────────────────────────────────────────

def index_all(dest: Path, index_name: str, model: str, namespace: str, limit: int):
    pc = get_pinecone()
    index = ensure_index(pc, index_name, model)
    meta_map = load_meta_map(dest)
    done = _load_ckpt()
    _enable_cost_meter()

    files = [p for p in dest.rglob("*")
             if p.is_file() and p.suffix.lower() in SUPPORTED_EXT and p.name != "archive_log.json"]
    if limit:
        files = files[:limit]
    if not files:
        print(f"\n📂 Tidak ada file terdukung di {dest}. Jalankan ingest_archive.py --apply dulu.\n")
        return

    from collections import Counter as _C
    _ext = _C(p.suffix.lower() for p in files)
    print(f"\n{'='*64}\n  KS Pinecone Indexer — {len(files)} file di {dest}")
    print(f"  tipe: {dict(_ext.most_common())}")
    print(f"  index='{index_name}' model='{model}' ns='{namespace}'\n{'='*64}\n")

    n_doc, n_skip, n_vec, batch = 0, 0, 0, []
    import extract_fields as EF

    def flush(force=False):
        nonlocal batch, n_vec
        while batch and (force or len(batch) >= BATCH):
            chunk, batch = batch[:BATCH], batch[BATCH:]
            retry(lambda c=chunk: index.upsert_records(namespace=namespace, records=c), what="upsert")
            n_vec += len(chunk)

    # File yg belum selesai (checkpoint) → diproses paralel.
    todo = []
    for p in files:
        if doc_id(str(p.relative_to(dest))) in done:
            n_skip += 1
        else:
            todo.append(p)
    print(f"  {n_skip:,} sudah ke-index (skip) | {len(todo):,} diproses | {WORKERS} worker paralel\n")

    def process_one(p):
        """Kerja berat per-file (OCR+extract) — dijalankan di worker thread."""
        base = doc_id(str(p.relative_to(dest)))
        meta = dict(meta_map.get(str(p.resolve()), {}))
        pages = extract_pages_any(p, cap=cap_for(meta))    # PDF/gambar/docx/xlsx; cap per doc-type
        if ENABLE_FIELD_EXTRACT:
            try:
                EF.enrich(meta, " ".join(t for _, t in pages), use_llm=EXTRACT_USE_LLM)
            except Exception:
                pass
        return base, records_for(p, base, meta, pages=pages), p.name, (meta.get("company") or "?")

    interrupted = False
    exe = ThreadPoolExecutor(max_workers=WORKERS)
    futs = {exe.submit(process_one, p): p for p in todo}
    try:
        for fut in as_completed(futs):          # main thread: upsert/checkpoint/print (serial, aman)
            try:
                base, recs, name, comp = fut.result()
            except Exception as e:
                print(f"  ⚠ gagal {futs[fut].name[:44]}: {type(e).__name__}")
                continue
            batch.extend(recs)
            flush()
            done.add(base)
            n_doc += 1
            print(f"  ✓ {name[:46]:46} [{len(recs)} chunk] {comp[:20]}")
            if n_doc % 25 == 0:
                _save_ckpt(done)
                proj = _USAGE["cost"] / n_doc * len(todo) if n_doc else 0.0
                print(f"  … {n_doc:,}/{len(todo):,} dok, {n_vec:,} chunk, "
                      f"💰 ${_USAGE['cost']:.2f} (proj ~${proj:.0f})")
    except KeyboardInterrupt:
        interrupted = True
        print("\n  ⏸ Dihentikan user — membatalkan antrian & menyimpan checkpoint...")
    finally:
        exe.shutdown(wait=False, cancel_futures=True)
        flush(force=True)
        _save_ckpt(done)              # SELALU simpan (normal / Ctrl-C / error) → resume aman

    if interrupted:
        print(f"  ▶ Progres tersimpan ({len(done)} dok). Jalankan lagi "
              f"'python index_to_pinecone.py' untuk lanjut dari sini.")
    print(f"\n{'='*64}")
    print(f"  ✅ {n_doc} dokumen di-index ({n_vec} chunk), {n_skip} skip")
    print(f"  💰 Biaya Claude: ${_USAGE['cost']:.2f}  "
          f"(in {_USAGE['in']:,} / out {_USAGE['out']:,} token)")
    if n_doc:
        print(f"     rata-rata ${_USAGE['cost']/n_doc:.4f}/dok")
    print(f"  🔢 Namespace '{namespace}' di index '{index_name}'")
    print(f"{'='*64}\n")


# ── Query helper (tes) ────────────────────────────────────────────────────────

def search(query: str, top_k: int = 5, index_name: str = INDEX_NAME,
           namespace: str = NAMESPACE, **filters):
    """Cari natural-language + filter metadata opsional.
       search("faktur pajak", company="PT Krakatau Shipyard", department="Finance")"""
    pc = get_pinecone()
    index = retry(lambda: pc.Index(index_name), what="resolve index")
    flt = {k: {"$eq": v} for k, v in filters.items() if v}
    q = {"inputs": {"text": query}, "top_k": top_k}
    if flt:
        q["filter"] = flt
    res = retry(lambda: index.search(
        namespace=namespace, query=q,
        fields=["doc_name", "company", "department", "project",
                "filename", "expire_date", "page_number", "chunk_text"]), what="search")

    def g(obj, key, default=None):           # tahan response object ATAU dict
        if obj is None:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    hits = g(g(res, "result"), "hits") or []
    print(f"\n🔍 '{query}'" + (f"  filter={flt}" if flt else "") + f"  → {len(hits)} hit")
    print("─" * 60)
    for h in hits:
        f = g(h, "fields") or {}
        score = g(h, "score", 0.0) or 0.0    # SDK Hit: property .score (wire _score)
        proj = f" | proyek: {f.get('project')}" if f.get("project") else ""
        exp  = f" | expire: {f.get('expire_date')}" if f.get("expire_date") else ""
        print(f"  [{score:.3f}] {f.get('doc_name','')}  (hal {f.get('page_number','?')})")
        print(f"          {f.get('company','')} / {f.get('department','')}{proj}")
        print(f"          {f.get('filename','')}{exp}")
        print(f"          “{(f.get('chunk_text','') or '')[:120]}…”\n")
    return res


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(description="Index Arsip_Rapih → Pinecone (integrated embedding)")
    ap.add_argument("--dest", default=str(DEST_DIR), help="folder hasil ingest (default D:\\Arsip_Rapih)")
    ap.add_argument("--index", default=INDEX_NAME)
    ap.add_argument("--model", default=EMBED_MODEL)
    ap.add_argument("--namespace", default=NAMESPACE)
    ap.add_argument("--limit", type=int, default=None, help="batasi jumlah PDF (uji coba)")
    ap.add_argument("--reset-checkpoint", action="store_true", help="abaikan checkpoint, index ulang semua")
    ap.add_argument("--workers", type=int, default=None, help=f"worker paralel (default {WORKERS})")
    ap.add_argument("--query", default=None, help="mode tes: cari kalimat ini lalu keluar")
    args = ap.parse_args()

    if args.query:
        search(args.query, index_name=args.index, namespace=args.namespace)
        return

    if args.workers:
        globals()["WORKERS"] = args.workers
    if args.reset_checkpoint and CHECKPOINT.exists():
        CHECKPOINT.unlink()
    index_all(Path(args.dest), args.index, args.model, args.namespace, args.limit)


if __name__ == "__main__":
    main()
